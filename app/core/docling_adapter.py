"""
Módulo adaptador para processamento de documentos.

Este módulo fornece uma camada de abstração para interagir com bibliotecas de processamento de documentos,
isolando o código da aplicação das dependências externas e facilitando testes e manutenção.
"""

import os
import tempfile
import io
import uuid
from pathlib import Path
from typing import Dict, List, Optional, Union, Any

# Bibliotecas para processamento de documentos
import docx
import PyPDF2
import pandas as pd
import openpyxl
import markdown
from PIL import Image

# Importar serviço de imagens
from app.services.image_service import ImageExtractor
from app.core.config import RESULTS_DIR


class DoclingAdapter:
    """
    Adaptador para processamento de documentos.

    Esta classe encapsula as interações com várias bibliotecas de processamento de documentos,
    fornecendo métodos simplificados para extração de conteúdo.
    """

    def __init__(self):
        """Inicializa o adaptador com configurações padrão."""
        # Inicializar o extrator de imagens
        try:
            self.image_extractor = ImageExtractor()
            print("ImageExtractor inicializado com sucesso")
        except Exception as e:
            print(f"Erro ao inicializar ImageExtractor: {str(e)}")
            self.image_extractor = None

    def process_document(
        self,
        file_path: Union[str, Path],
        extract_text: bool = True,
        extract_tables: bool = True,
        extract_images: bool = False,
        extract_pages_as_images: bool = False,
    ) -> Dict[str, Any]:
        """
        Processa um documento usando bibliotecas específicas para cada tipo de arquivo.

        Args:
            file_path: Caminho para o arquivo a ser processado
            extract_text: Se deve extrair texto do documento
            extract_tables: Se deve extrair tabelas do documento
            extract_images: Se deve extrair imagens incorporadas do documento
            extract_pages_as_images: Se deve converter páginas inteiras em imagens (apenas para PDF)

        Returns:
            Dicionário com os resultados do processamento
        """
        try:
            file_path = str(file_path)  # Converter Path para string se necessário
            file_extension = os.path.splitext(file_path)[1].lower()

            # Inicializar resultado
            processing_result = {
                "status": "success",
                "message": "Documento processado com sucesso",
                "content": {},
            }

            # Processar com base na extensão do arquivo
            if file_extension == ".pdf":
                self._process_pdf(
                    file_path, processing_result, extract_text, extract_tables, extract_images, extract_pages_as_images
                )
            elif file_extension == ".docx":
                self._process_docx(
                    file_path, processing_result, extract_text, extract_tables, extract_images
                )
            elif file_extension in [".xlsx", ".xls"]:
                self._process_excel(file_path, processing_result, extract_text, extract_tables)
            else:
                processing_result["status"] = "error"
                processing_result["message"] = f"Formato de arquivo não suportado: {file_extension}"
                return processing_result

            return processing_result

        except Exception as e:
            # Registrar erro e retornar informações sobre a falha
            return {
                "status": "error",
                "message": f"Erro ao processar documento: {str(e)}",
                "content": None,
            }

    def _process_pdf(self, file_path, result, extract_text, extract_tables, extract_images, extract_pages_as_images=False):
        """
        Processa um arquivo PDF.

        Args:
            file_path: Caminho para o arquivo PDF
            result: Dicionário para armazenar os resultados
            extract_text: Se deve extrair texto
            extract_tables: Se deve extrair tabelas
            extract_images: Se deve extrair imagens incorporadas
            extract_pages_as_images: Se deve converter páginas inteiras em imagens
        """
        with open(file_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)

            # Extrair texto
            if extract_text:
                text = ""
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n\n"

                result["content"]["text"] = text
                result["content"]["markdown"] = text  # Texto simples como markdown
                result["content"]["html"] = f"<pre>{text}</pre>"  # Texto simples como HTML

            # Extrair imagens
            if extract_images:
                print("Iniciando extração de imagens...")

                # Verificar se result é None
                if result is None:
                    print("ERRO: result é None!")
                    return

                # Verificar se content existe
                if "content" not in result:
                    print("ERRO: 'content' não existe em result!")
                    result["content"] = {}

                document_id = result.get("id", str(uuid.uuid4()))
                print(f"ID do documento: {document_id}")

                # Garantir que o diretório de imagens exista
                images_dir = os.path.join(RESULTS_DIR, document_id, "images")
                print(f"Diretório de imagens: {images_dir}")
                os.makedirs(images_dir, exist_ok=True)

                # Garantir que metadata existe
                if "metadata" not in result:
                    print("Criando dicionário metadata")
                    result["metadata"] = {}

                # Verificar se o extrator de imagens está inicializado
                if self.image_extractor is None:
                    print("Erro: ImageExtractor não está inicializado")
                    result["metadata"]["image_extraction_error"] = "Extrator de imagens não inicializado"
                    return

                try:
                    print(f"Chamando extract_from_pdf com extract_pages={extract_pages_as_images}")
                    images_result = self.image_extractor.extract_from_pdf(
                        file_path=file_path,
                        images_dir=images_dir,
                        extract_pages=extract_pages_as_images
                    )

                    print(f"Resultado da extração: {images_result}")

                    if images_result.get("success"):
                        print("Extração bem-sucedida")
                        result["content"]["images"] = images_result.get("images", [])
                        result["metadata"]["image_count"] = len(images_result.get("images", []))
                        print(f"Número de imagens extraídas: {result['metadata']['image_count']}")
                    else:
                        print(f"Extração falhou: {images_result.get('error', 'Erro desconhecido')}")
                        result["metadata"]["image_extraction_error"] = images_result.get("error", "Erro desconhecido")
                except Exception as e:
                    print(f"Erro ao extrair imagens: {str(e)}")
                    print(f"Tipo de erro: {type(e)}")
                    print(f"Erro detalhado: {repr(e)}")
                    import traceback
                    traceback.print_exc()
                    result["metadata"]["image_extraction_error"] = f"Erro ao extrair imagens: {str(e)}"

            # Metadados
            if "metadata" not in result:
                result["metadata"] = {}

            result["metadata"].update({
                "pages": len(pdf_reader.pages),
                "title": (
                    pdf_reader.metadata.title
                    if pdf_reader.metadata and hasattr(pdf_reader.metadata, "title")
                    else "Sem título"
                ),
            })

    def _process_docx(self, file_path, result, extract_text, extract_tables, extract_images):
        """Processa um arquivo DOCX."""
        doc = docx.Document(file_path)

        # Extrair texto
        if extract_text:
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            result["content"]["text"] = text

            # Converter para markdown (texto simples com formatação básica)
            md_text = ""
            for para in doc.paragraphs:
                if para.style.name.startswith("Heading"):
                    level = int(para.style.name.replace("Heading", ""))
                    md_text += "#" * level + " " + para.text + "\n\n"
                else:
                    md_text += para.text + "\n\n"

            result["content"]["markdown"] = md_text
            result["content"]["html"] = markdown.markdown(md_text)

        # Extrair tabelas
        if extract_tables and doc.tables:
            tables = []
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append({"page": 1, "data": table_data})  # DOCX não tem conceito de página

            result["content"]["tables"] = tables

        # Extrair imagens
        if extract_images:
            document_id = result.get("id", str(uuid.uuid4()))

            # Garantir que o diretório de imagens exista
            images_dir = os.path.join(RESULTS_DIR, document_id, "images")
            os.makedirs(images_dir, exist_ok=True)

            # Garantir que metadata existe
            if "metadata" not in result:
                result["metadata"] = {}

            # Verificar se o extrator de imagens está inicializado
            if self.image_extractor is None:
                print("Erro: ImageExtractor não está inicializado")
                result["metadata"]["image_extraction_error"] = "Extrator de imagens não inicializado"
                return

            try:
                images_result = self.image_extractor.extract_from_docx(
                    file_path=file_path,
                    images_dir=images_dir,
                    extract_pages=False
                )

                if images_result.get("success"):
                    result["content"]["images"] = images_result.get("images", [])
                    result["metadata"]["image_count"] = len(images_result.get("images", []))
                else:
                    result["metadata"]["image_extraction_error"] = images_result.get("error", "Erro desconhecido")
            except Exception as e:
                print(f"Erro ao extrair imagens: {str(e)}")
                result["metadata"]["image_extraction_error"] = f"Erro ao extrair imagens: {str(e)}"

        # Metadados
        if "metadata" not in result:
            result["metadata"] = {}

        result["metadata"].update({
            "title": (
                doc.core_properties.title
                if hasattr(doc, "core_properties") and hasattr(doc.core_properties, "title")
                else "Sem título"
            ),
            "pages": 1,  # DOCX não tem conceito de página
        })

    def _process_excel(self, file_path, result, extract_text, extract_tables):
        """Processa um arquivo Excel."""
        import simplejson as json
        import numpy as np

        # Usar pandas para ler o Excel
        excel_file = pd.ExcelFile(file_path)
        sheet_names = excel_file.sheet_names

        # Extrair tabelas
        if extract_tables:
            tables = []
            for sheet_name in sheet_names:
                try:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)

                    # Substituir NaN por None para compatibilidade com JSON
                    df = df.replace({np.nan: None})

                    # Extrair cabeçalhos e dados
                    headers = df.columns.tolist()

                    # Converter cada linha para lista, substituindo NaN por None
                    data = []
                    for _, row in df.iterrows():
                        row_data = []
                        for col in headers:
                            val = row[col]
                            if pd.isna(val):
                                row_data.append(None)
                            else:
                                row_data.append(val)
                        data.append(row_data)

                    tables.append({
                        "page": 1,  # Excel não tem conceito de página
                        "sheet": sheet_name,
                        "data": data,
                        "headers": headers,
                    })
                except Exception as e:
                    # Em caso de erro, adicionar informação de erro
                    tables.append({
                        "page": 1,
                        "sheet": sheet_name,
                        "error": str(e),
                        "headers": [],
                        "data": []
                    })

            result["content"]["tables"] = tables

        # Extrair texto (cabeçalhos e dados como texto)
        if extract_text:
            text = ""
            for sheet_name in sheet_names:
                try:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    # Substituir NaN por strings vazias para exibição de texto
                    df = df.fillna("")
                    text += f"Sheet: {sheet_name}\n\n"
                    text += df.to_string(index=False, na_rep="") + "\n\n"
                except Exception as e:
                    text += f"Sheet: {sheet_name}\n\nErro ao processar: {str(e)}\n\n"

            result["content"]["text"] = text
            result["content"]["markdown"] = text  # Texto simples como markdown
            result["content"]["html"] = f"<pre>{text}</pre>"  # Texto simples como HTML

        # Metadados
        result["metadata"] = {"title": os.path.basename(file_path), "sheets": sheet_names}

    def get_document_metadata(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Extrai metadados de um documento.

        Args:
            file_path: Caminho para o arquivo

        Returns:
            Dicionário com metadados do documento
        """
        try:
            file_path = str(file_path)  # Converter Path para string se necessário
            file_extension = os.path.splitext(file_path)[1].lower()

            metadata = {
                "title": os.path.basename(file_path),
                "format": file_extension[1:] if file_extension.startswith(".") else file_extension,
            }

            # Extrair metadados específicos por tipo de arquivo
            if file_extension == ".pdf":
                with open(file_path, "rb") as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata["pages"] = len(pdf_reader.pages)
                    if (
                        pdf_reader.metadata
                        and hasattr(pdf_reader.metadata, "title")
                        and pdf_reader.metadata.title
                    ):
                        metadata["title"] = pdf_reader.metadata.title

            elif file_extension == ".docx":
                doc = docx.Document(file_path)
                if (
                    hasattr(doc, "core_properties")
                    and hasattr(doc.core_properties, "title")
                    and doc.core_properties.title
                ):
                    metadata["title"] = doc.core_properties.title

            elif file_extension in [".xlsx", ".xls"]:
                excel_file = pd.ExcelFile(file_path)
                metadata["sheets"] = excel_file.sheet_names

            return metadata

        except Exception as e:
            return {"status": "error", "message": f"Erro ao extrair metadados: {str(e)}"}

    def convert_to_format(
        self, file_path: Union[str, Path], output_format: str = "markdown"
    ) -> Optional[str]:
        """
        Converte um documento para o formato especificado.

        Args:
            file_path: Caminho para o arquivo a ser convertido
            output_format: Formato de saída desejado (markdown, html, text)

        Returns:
            Conteúdo convertido ou None em caso de erro
        """
        try:
            # Processar o documento para extrair o conteúdo
            result = self.process_document(file_path, extract_text=True)

            if result["status"] != "success" or not result.get("content"):
                return None

            # Retornar o formato solicitado
            if output_format.lower() == "markdown" and "markdown" in result["content"]:
                return result["content"]["markdown"]
            elif output_format.lower() == "html" and "html" in result["content"]:
                return result["content"]["html"]
            elif output_format.lower() == "text" and "text" in result["content"]:
                return result["content"]["text"]
            else:
                raise ValueError(f"Formato de saída não suportado: {output_format}")

        except Exception as e:
            print(f"Erro ao converter documento: {str(e)}")
            return None
